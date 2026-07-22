import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from llm.schema import ExtractedDocument

SCALAR_FIELDS = [
    "invoice_number",
    "invoice_date",
    "due_date",
    "vendor_name",
    "vendor_address",
    "customer_name",
    "subtotal",
    "tax",
    "total",
    "currency",
]

LINE_ITEM_HEADERS = ["description", "quantity", "unit_price", "amount"]


def _add_table(document: Document, rows: list[list]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        cells = table.rows[r].cells
        for c in range(col_count):
            cells[c].text = str(row[c]) if c < len(row) else ""
            if r == 0:
                for paragraph in cells[c].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def build_word_report(data: ExtractedDocument) -> bytes:
    document = Document()

    title = document.add_heading(data.document_type or "Extracted Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if data.summary:
        document.add_paragraph(data.summary)

    detail_rows = [
        [field.replace("_", " ").title(), str(getattr(data, field))]
        for field in SCALAR_FIELDS
        if getattr(data, field) is not None
    ]
    for extra in data.fields:
        detail_rows.append([extra.label, extra.value])

    if detail_rows:
        document.add_heading("Details", level=2)
        _add_table(document, [["Field", "Value"]] + detail_rows)

    if data.line_items:
        document.add_heading("Line Items", level=2)
        rows = [LINE_ITEM_HEADERS] + [
            [str(getattr(item, header)) for header in LINE_ITEM_HEADERS] for item in data.line_items
        ]
        _add_table(document, rows)

    for table in data.tables:
        rows = ([table.headers] if table.headers else []) + table.rows
        if not rows:
            continue
        document.add_heading(table.title or "Table", level=2)
        _add_table(document, rows)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
